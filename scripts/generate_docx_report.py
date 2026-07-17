"""Generate DOCX report for Graph ML project — ~20 pages, English, academic style.

Usage:
    python scripts/generate_docx_report.py
"""
from __future__ import annotations

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

PROJECT = "/Users/ducnguyen/Projects/HUST2/IT5429E/GraphMl"
FIGURES = os.path.join(PROJECT, "figures")
OUTPUT = os.path.join(PROJECT, "report", "graphml-report.docx")


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None, highlight_last=False):
    """Add a formatted table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
        set_cell_shading(cell, "2C3E50")
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        is_highlight = highlight_last and r_idx == len(rows) - 1
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)
                if is_highlight:
                    run.bold = True
        if is_highlight:
            for cell in row.cells:
                set_cell_shading(cell, "E8F8F5")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def add_figure(doc, img_path, caption, width=Inches(5.5)):
    """Add a centered figure with caption."""
    if not os.path.exists(img_path):
        doc.add_paragraph(f"[Figure not found: {img_path}]")
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=width)

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = doc.styles['Caption'] if 'Caption' in [s.name for s in doc.styles] else doc.styles['Normal']
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(10)


def build_report():
    doc = Document()

    # --- Page setup ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- Default font ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # ================================================================
    # TITLE PAGE
    # ================================================================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Residual Multi-Head Mixing GAT\nfor Semi-Supervised Node Classification")
    run.bold = True
    run.font.size = Pt(24)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("An Experimental Study on the Cora Citation Network")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(3):
        doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Graph Machine Learning — Course Project\n")
    run.font.size = Pt(12)
    run = info.add_run("Hanoi University of Science and Technology\n\n")
    run.font.size = Pt(12)
    run = info.add_run("Team Members:\n")
    run.font.size = Pt(12)
    run.bold = True

    members = [
        "Do Tuan Minh (20261057M)",
        "Hoang Huy Chien (20261069M)",
        "Tran Tien Dung (20252574M)",
        "Tran Manh Tien (20252762M)",
        "Nguyen Tien Duc (20252076M)",
    ]
    for m in members:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(m)
        run.font.size = Pt(11)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("July 2026")
    run.font.size = Pt(12)

    doc.add_page_break()

    # ================================================================
    # ABSTRACT
    # ================================================================
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "This project studies semi-supervised node classification on the Cora citation network "
        "using Graph Neural Networks (GNNs). We compare a feature-only Multi-Layer Perceptron (MLP) "
        "baseline against three established GNN architectures — Graph Convolutional Network (GCN), "
        "GraphSAGE, and Graph Attention Network (GAT) — and propose a Residual Multi-Head Mixing GAT "
        "that combines GATv2 multi-head attention with an input-feature skip connection and DropEdge "
        "regularization. Experiments over 5 random seeds demonstrate that graph structure improves "
        "accuracy by +24% over MLP (0.5888 vs 0.82-0.83), and the proposed model achieves the best "
        "test accuracy of 0.8326, outperforming GAT by +0.5%. Ablation analysis reveals that the "
        "residual skip connection contributes +1.7% improvement, while DropEdge has negligible effect "
        "on this small graph. We discuss implications, limitations, and directions for future work."
    )

    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS (placeholder)
    # ================================================================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ("1. Introduction", "3"),
        ("2. Background", "4"),
        ("   2.1 Graph Neural Networks", "4"),
        ("   2.2 Graph Attention Networks", "5"),
        ("   2.3 The Cora Dataset", "6"),
        ("3. Related Work", "7"),
        ("4. Proposed Method", "8"),
        ("   4.1 Architecture Overview", "8"),
        ("   4.2 GATv2 Attention", "9"),
        ("   4.3 Residual Skip Connection", "9"),
        ("   4.4 DropEdge Regularization", "10"),
        ("5. Experiments", "11"),
        ("   5.1 Experimental Setup", "11"),
        ("   5.2 Baseline Comparison: MLP vs GCN", "12"),
        ("   5.3 Architecture Comparison", "13"),
        ("   5.4 Ablation Study", "14"),
        ("   5.5 Confusion Matrix Analysis", "15"),
        ("6. Discussion", "16"),
        ("7. Limitations and Future Work", "18"),
        ("8. Conclusion", "19"),
        ("References", "20"),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{item}")
        run.font.size = Pt(11)
        tab = p.add_run(f"\t{page}")
        tab.font.size = Pt(11)
        tab.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ================================================================
    # 1. INTRODUCTION
    # ================================================================
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "Graph-structured data is ubiquitous in real-world applications: social networks, "
        "citation networks, molecular graphs, and knowledge graphs all exhibit relational "
        "structure that traditional machine learning models cannot fully exploit. Graph Neural "
        "Networks (GNNs) have emerged as a powerful framework for learning on such data by "
        "propagating information along edges through a mechanism known as message passing."
    )
    doc.add_paragraph(
        "In this project, we focus on semi-supervised node classification — the task of "
        "predicting labels for unlabeled nodes given a small set of labeled examples and the "
        "graph structure. We use the Cora citation network as our benchmark dataset, where "
        "nodes represent scientific papers, edges represent citations, and the goal is to "
        "classify each paper into one of seven research topics."
    )
    doc.add_paragraph(
        "Our contributions are threefold:"
    )
    contribs = [
        "We implement and compare four models: MLP (baseline), GCN, GraphSAGE, and GAT, "
        "establishing a comprehensive benchmark on Cora.",
        "We propose a Residual Multi-Head Mixing GAT that augments GATv2 with an "
        "input-feature skip connection (H₀·W_skip) and DropEdge regularization.",
        "We conduct a thorough ablation study to isolate the contribution of each "
        "component: residual connection (+1.7%), DropEdge (negligible), and attention heads.",
    ]
    for i, c in enumerate(contribs, 1):
        doc.add_paragraph(f"{i}. {c}", style='List Number')

    doc.add_paragraph(
        "The remainder of this report is organized as follows: Section 2 provides background "
        "on GNNs and the Cora dataset. Section 3 reviews related work. Section 4 describes "
        "our proposed method. Section 5 presents experiments and results. Section 6 discusses "
        "the findings. Section 7 addresses limitations and future directions. Section 8 concludes."
    )

    # ================================================================
    # 2. BACKGROUND
    # ================================================================
    doc.add_heading('2. Background', level=1)

    doc.add_heading('2.1 Graph Neural Networks', level=2)
    doc.add_paragraph(
        "Graph Neural Networks generalize neural networks to graph-structured data. The core "
        "idea is message passing: at each layer, every node aggregates feature information from "
        "its neighbors and combines it with its own features to produce an updated representation. "
        "Formally, for a node v at layer l:"
    )
    doc.add_paragraph(
        "    hᵥ⁽ˡ⁾ = COMBINE⁽ˡ⁾(hᵥ⁽ˡ⁻¹⁾, AGGREGATE⁽ˡ⁾({hᵤ⁽ˡ⁻¹⁾ : u ∈ N(v)}))"
    )
    doc.add_paragraph(
        "Different GNN architectures vary in how they define the AGGREGATE and COMBINE functions. "
        "GCN uses a spectral convolution based on the normalized adjacency matrix. GraphSAGE "
        "uses sampling and various aggregation functions (mean, LSTM, pool). GAT introduces "
        "learned attention weights to assign different importance to different neighbors."
    )

    doc.add_heading('2.2 Graph Attention Networks (GAT)', level=2)
    doc.add_paragraph(
        "GAT (Veličković et al., 2018) computes attention coefficients between connected nodes "
        "using a shared linear transformation followed by a LeakyReLU activation and softmax "
        "normalization. Multi-head attention is used to stabilize the learning process by "
        "concatenating K independent attention mechanisms. GATv2 (Brody et al., 2022) fixes "
        "a static ranking problem in the original GAT by modifying the order of operations, "
        "resulting in truly dynamic attention coefficients."
    )
    doc.add_paragraph(
        "In our proposed model, we use GATv2 as the base attention mechanism, with 8 attention "
        "heads in the first layer (concatenated to 64 dimensions) and a single attention head "
        "in the second layer projecting to the 7 output classes."
    )

    doc.add_heading('2.3 The Cora Dataset', level=2)
    doc.add_paragraph(
        "Cora is a widely-used citation network dataset for semi-supervised node classification. "
        "It consists of 2,708 scientific publications classified into one of seven research "
        "topics. The citation links form a directed graph with 5,429 edges (after making "
        "edges undirected: 10,556). Each paper is represented by a 1,433-dimensional "
        "binary bag-of-words feature vector indicating the presence of specific words."
    )

    # Dataset statistics table
    doc.add_paragraph()
    add_styled_table(doc,
        ["Property", "Value"],
        [
            ["Nodes (papers)", "2,708"],
            ["Edges (citations)", "5,278 directed / 10,556 undirected"],
            ["Features", "1,433 (binary bag-of-words)"],
            ["Classes", "7 (research topics)"],
            ["Train / Validation / Test", "140 / 500 / 1,000"],
            ["Graph density", "~0.14% (very sparse)"],
            ["Largest connected component", "2,485 nodes"],
        ],
        col_widths=[2.0, 4.0]
    )
    doc.add_paragraph("Table 1: Cora dataset statistics.", style='Caption')

    doc.add_paragraph(
        "The standard Planetoid split (Yang et al., 2016) is used throughout: 140 nodes for "
        "training (20 per class), 500 for validation, and 1,000 for testing. Features are "
        "row-normalized using PyG's NormalizeFeatures transform. This split is fixed across "
        "all experiments to ensure fair comparison."
    )

    # ================================================================
    # 3. RELATED WORK
    # ================================================================
    doc.add_heading('3. Related Work', level=1)
    doc.add_paragraph(
        "Graph Neural Networks have seen rapid development since the introduction of "
        "spectral methods (Bruna et al., 2014) and their simplification into GCN "
        "(Kipf & Welling, 2017). GCN's success on semi-supervised classification "
        "sparked numerous extensions:"
    )

    related = [
        ("GraphSAGE (Hamilton et al., 2017)", "Introduced inductive learning on graphs "
         "through sampling and aggregation, enabling generalization to unseen nodes."),
        ("GAT (Veličković et al., 2018)", "Applied attention mechanisms to graph "
         "neighborhoods, learning to weight neighbors differently."),
        ("GATv2 (Brody et al., 2022)", "Fixed the static attention ranking problem in "
         "GAT, achieving truly dynamic attention."),
        ("DropEdge (Rong et al., 2020)", "Randomly removes edges during training as "
         "data augmentation to mitigate over-smoothing in deep GNNs."),
        ("Residual connections (He et al., 2016)", "Skip connections from ResNet that "
         "enable gradient flow in deep networks, adapted for GNNs."),
    ]
    for name, desc in related:
        p = doc.add_paragraph()
        run = p.add_run(f"• {name}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "Our work combines GATv2 attention, residual skip connections from raw features, "
        "and DropEdge regularization into a unified architecture. While each component "
        "has been studied individually, their combination for shallow (2-layer) GNNs on "
        "small citation networks has not been thoroughly investigated."
    )

    # ================================================================
    # 4. PROPOSED METHOD
    # ================================================================
    doc.add_heading('4. Proposed Method', level=1)

    doc.add_heading('4.1 Architecture Overview', level=2)
    doc.add_paragraph(
        "We propose the Residual Multi-Head Mixing GAT, a two-layer graph attention "
        "network augmented with an input-feature skip connection and DropEdge "
        "regularization. The architecture processes input features through the "
        "following pipeline:"
    )
    doc.add_paragraph(
        "    H₀ = X                                        # Raw features (1433-dim)\n"
        "    H₁ = ELU(GATv2_multi(H₀, E))         # 8 heads, concat → 64-dim\n"
        "    H₂ = GATv2_single(H₁, E)                # Single head → 7-dim\n"
        "    Z  = H₂ + H₀ · W_skip                 # Residual → logits"
    )
    doc.add_paragraph(
        "The model takes raw features X (1,433 dimensions) and the edge index E as input. "
        "The first GATv2 layer applies 8 attention heads with 8 hidden dimensions each, "
        "concatenated to produce H₁ (64 dimensions) with ELU activation. The second GATv2 "
        "layer uses a single attention head to project H₁ to the 7 output classes, producing H₂. "
        "Finally, a residual skip connection adds a linear projection of the raw input features "
        "(W_skip ∈ ℝ^{1433×7}, no bias) to produce the output logits Z."
    )

    doc.add_heading('4.2 GATv2 Attention', level=2)
    doc.add_paragraph(
        "We use GATv2 instead of the original GAT to avoid the static attention ranking problem. "
        "In GAT, the attention score is computed as e_ij = LeakyReLU(a^T [Wh_i || Wh_j]), which "
        "can be shown to produce a static ranking of neighbors independent of the query node. "
        "GATv2 modifies this to e_ij = a^T LeakyReLU([Wh_i || Wh_j]), where the linear "
        "transformation is applied after concatenation, resulting in truly dynamic attention "
        "that depends on both the query and key nodes."
    )
    doc.add_paragraph(
        "The first layer uses K=8 attention heads, each computing 8-dimensional representations. "
        "These are concatenated to produce a 64-dimensional hidden state. The second layer uses "
        "a single head to produce the final 7-dimensional output (one logit per class)."
    )

    doc.add_heading('4.3 Residual Skip Connection', level=2)
    doc.add_paragraph(
        "The key innovation of our model is the residual skip connection from raw input "
        "features directly to the output logits:"
    )
    doc.add_paragraph("    Z = H₂ + W_skip · H₀")
    doc.add_paragraph(
        "where W_skip ∈ ℝ^{7×1433} is a learnable linear projection (without bias). "
        "This skip connection serves multiple purposes:"
    )
    purposes = [
        "Gradient flow: It provides a direct path for gradients from the loss to the input "
        "features, mitigating potential vanishing gradient issues even in shallow networks.",
        "Over-smoothing mitigation: By preserving raw feature information at the output, "
        "the model can counteract the over-smoothing effect where node representations "
        "become indistinguishable after multiple message-passing steps.",
        "Capacity augmentation: The linear projection W_skip acts as a simple linear "
        "classifier on raw features, complementing the graph-based predictions from the GAT layers.",
    ]
    for p_text in purposes:
        doc.add_paragraph(p_text, style='List Bullet')

    doc.add_heading('4.4 DropEdge Regularization', level=2)
    doc.add_paragraph(
        "DropEdge (Rong et al., 2020) randomly removes a fraction p of edges from the "
        "graph during each training step. In our model, we apply DropEdge with p=0.2, "
        "meaning 20% of edges are randomly dropped during training. This serves as "
        "structural data augmentation and regularization:"
    )
    doc.add_paragraph(
        "• During training: edges are randomly dropped before each forward pass\n"
        "• During evaluation: all edges are used (no dropout)\n"
        "• Effect: forces the model to not over-rely on any specific edge, promoting robustness"
    )
    doc.add_paragraph(
        "The ablation switch allows us to independently control the residual connection "
        "and DropEdge, enabling us to isolate their individual contributions."
    )

    # ================================================================
    # 5. EXPERIMENTS
    # ================================================================
    doc.add_heading('5. Experiments', level=1)

    doc.add_heading('5.1 Experimental Setup', level=2)
    doc.add_paragraph(
        "All models share the same training protocol to ensure fair comparison. "
        "We use the Adam optimizer with a learning rate of 0.005 and weight decay of 5×10⁻⁴. "
        "Models are trained for up to 200 epochs with early stopping (patience=100 on "
        "validation accuracy). The model checkpoint at the best validation epoch is used "
        "to report the final test accuracy."
    )

    add_styled_table(doc,
        ["Hyperparameter", "Value"],
        [
            ["Optimizer", "Adam"],
            ["Learning rate", "0.005"],
            ["Weight decay", "5 × 10⁻⁴"],
            ["Max epochs", "200"],
            ["Early stopping patience", "100"],
            ["Dropout rate", "0.6"],
            ["Hidden dimension (GCN/SAGE)", "64"],
            ["GAT per-head dimension", "8"],
            ["Attention heads", "8"],
            ["DropEdge probability", "0.2 (proposed only)"],
        ],
        col_widths=[2.5, 2.5]
    )
    doc.add_paragraph("Table 2: Hyperparameters used across all experiments.", style='Caption')

    doc.add_paragraph(
        "For reproducibility, each experiment is run with 5 different random seeds "
        "(42, 43, 44, 45, 46). We report the mean and standard deviation of test accuracy "
        "across seeds. All random number generators (Python, NumPy, PyTorch, cuDNN) are "
        "seeded deterministically at the start of each run."
    )
    doc.add_paragraph(
        "The implementation uses PyTorch 2.13 and PyTorch Geometric (PyG) 2.8. "
        "All experiments are run on CPU as the Cora dataset is small enough that GPU "
        "acceleration provides no significant speedup."
    )

    # --- 5.2 Baselines ---
    doc.add_heading('5.2 Baseline Comparison: MLP vs GCN', level=2)
    doc.add_paragraph(
        "We first establish whether graph structure provides value by comparing a "
        "feature-only MLP baseline against a 2-layer GCN. The MLP uses two fully "
        "connected layers with ReLU activation and dropout, completely ignoring the "
        "edge structure. The GCN uses spectral convolution with the normalized "
        "adjacency matrix."
    )

    add_styled_table(doc,
        ["Model", "Test Acc (mean ± std)", "Val Acc (mean)", "Seeds"],
        [
            ["MLP", "0.5888 ± 0.0104", "0.6152", "5"],
            ["GCN", "0.8276 ± 0.0050", "0.8156", "5"],
        ],
        col_widths=[1.2, 2.0, 1.5, 0.8]
    )
    doc.add_paragraph("Table 3: MLP vs GCN baseline results.", style='Caption')

    doc.add_paragraph(
        "The results are striking: GCN achieves 0.8276 test accuracy compared to MLP's "
        "0.5888, a relative improvement of +40.6% (absolute +23.9%). This confirms that "
        "the citation graph structure carries crucial information for paper classification "
        "that cannot be captured by bag-of-words features alone. Papers in the same research "
        "field tend to cite each other, creating clusters in the graph that GNNs can exploit "
        "through message passing."
    )
    doc.add_paragraph(
        "The low standard deviation for GCN (0.005) also indicates stable training "
        "across different random initializations, while MLP shows slightly higher "
        "variance (0.010), suggesting that graph-based learning is more robust."
    )

    add_figure(doc,
        os.path.join(FIGURES, "baselines_curves.png"),
        "Figure 1: Training loss and validation accuracy curves for MLP vs GCN."
    )

    # --- 5.3 Architecture Comparison ---
    doc.add_heading('5.3 Architecture Comparison: GCN vs GraphSAGE vs GAT', level=2)
    doc.add_paragraph(
        "We next compare three GNN architectures to understand how different message-passing "
        "mechanisms affect performance on Cora."
    )

    add_styled_table(doc,
        ["Model", "Test Acc (mean ± std)", "Val Acc (mean)", "Seeds"],
        [
            ["GraphSAGE", "0.8130 ± 0.0077", "0.8000", "5"],
            ["GAT", "0.8272 ± 0.0080", "0.8160", "5"],
            ["GCN", "0.8276 ± 0.0050", "0.8156", "5"],
        ],
        col_widths=[1.2, 2.0, 1.5, 0.8]
    )
    doc.add_paragraph("Table 4: Architecture comparison results.", style='Caption')

    doc.add_paragraph(
        "GCN (0.8276) and GAT (0.8272) achieve nearly identical performance, with GCN "
        "marginally ahead by 0.04%. This is surprising given that GAT's attention mechanism "
        "should theoretically provide more expressive power. On the small, homogeneous Cora "
        "dataset, simple mean aggregation (GCN) is already sufficient, and the additional "
        "complexity of attention does not translate to meaningful improvements."
    )
    doc.add_paragraph(
        "GraphSAGE (0.8130) underperforms both GCN and GAT by approximately 1.5%. "
        "This may be due to its mean aggregator being less effective than GCN's spectral "
        "convolution on sparse citation graphs, or the absence of the normalized adjacency "
        "preprocessing that GCN applies."
    )

    add_figure(doc,
        os.path.join(FIGURES, "model_comparison_curves.png"),
        "Figure 2: Training loss and validation accuracy curves for all five models."
    )

    add_figure(doc,
        os.path.join(FIGURES, "accuracy_bar_chart.png"),
        "Figure 3: Test accuracy comparison across all models with error bars (±1 std)."
    )

    # --- 5.4 Ablation Study ---
    doc.add_heading('5.4 Ablation Study', level=2)
    doc.add_paragraph(
        "To understand the contribution of each component in our proposed model, we conduct "
        "an ablation study that systematically enables or disables the residual skip "
        "connection, DropEdge, and varies the number of attention heads."
    )

    add_styled_table(doc,
        ["Variant", "Residual", "DropEdge", "Heads", "Test Acc (mean ± std)", "Δ vs Base"],
        [
            ["GAT (base)", "No", "0.0", "8", "0.8164 ± 0.0130", "—"],
            ["+ residual", "Yes", "0.0", "8", "0.8334 ± 0.0071", "+1.7%"],
            ["+ res + dropedge", "Yes", "0.2", "8", "0.8326 ± 0.0080", "+1.6%"],
            ["full, heads=4", "Yes", "0.2", "4", "0.8260 ± 0.0139", "+1.0%"],
        ],
        col_widths=[1.4, 0.7, 0.7, 0.5, 1.8, 0.8]
    )
    doc.add_paragraph("Table 5: Ablation study results for the proposed model.", style='Caption')

    doc.add_paragraph(
        "Residual Skip Connection (+1.7%): Adding the residual connection H₀·W_skip "
        "improves test accuracy from 0.8164 to 0.8334, the largest single-component "
        "improvement. The skip connection provides a direct path for raw feature information "
        "to reach the output, complementing the graph-convolved representations. It also "
        "facilitates gradient flow and may help mitigate over-smoothing even in our shallow "
        "2-layer architecture."
    )
    doc.add_paragraph(
        "DropEdge (−0.1%): Adding DropEdge with p=0.2 has a negligible negative effect "
        "(0.8334 → 0.8326). We hypothesize that Cora's small size (2,708 nodes, 5,278 edges) "
        "means the model does not overfit severely enough for regularization to help. "
        "Additionally, our 2-layer architecture is shallow enough that over-smoothing is "
        "not a significant concern. DropEdge is expected to be more beneficial on larger "
        "graphs or deeper architectures."
    )
    doc.add_paragraph(
        "Reduced Heads (−0.7%): Reducing attention heads from 8 to 4 decreases accuracy "
        "from 0.8326 to 0.8260. This confirms that multi-head attention provides value by "
        "allowing the model to attend to different aspects of the neighborhood simultaneously. "
        "However, the improvement is moderate, suggesting diminishing returns beyond 4 heads."
    )

    add_figure(doc,
        os.path.join(FIGURES, "ablation_bar_chart.png"),
        "Figure 4: Ablation study — test accuracy for each variant with error bars."
    )

    # --- 5.5 Confusion Matrix ---
    doc.add_heading('5.5 Confusion Matrix Analysis', level=2)
    doc.add_paragraph(
        "To understand per-class performance, we visualize the confusion matrix of the "
        "proposed model on the test set (1,000 nodes). The confusion matrix reveals "
        "which classes are most frequently confused with each other."
    )

    add_figure(doc,
        os.path.join(FIGURES, "confusion_matrix.png"),
        "Figure 5: Confusion matrix of the proposed model on the Cora test set."
    )

    doc.add_paragraph(
        "The diagonal entries (correct predictions) are dominant for most classes, "
        "indicating strong classification performance. Misclassifications tend to occur "
        "between semantically related research topics (e.g., neural networks and machine "
        "learning), which is expected given the overlap in vocabulary between these fields."
    )

    # Val-test scatter
    add_figure(doc,
        os.path.join(FIGURES, "val_vs_test_scatter.png"),
        "Figure 6: Validation vs test accuracy scatter plot across all models and seeds."
    )

    doc.add_paragraph(
        "The validation-test scatter plot shows a consistent positive correlation between "
        "validation and test accuracy (r ≈ 0.6), confirming that our validation-based "
        "model selection strategy is effective. The val-test gap of approximately 1-2% "
        "indicates mild overfitting that is well-controlled by our training protocol."
    )

    # ================================================================
    # 6. DISCUSSION
    # ================================================================
    doc.add_heading('6. Discussion', level=1)

    doc.add_heading('6.1 The Importance of Graph Structure', level=2)
    doc.add_paragraph(
        "Our most significant finding is the dramatic improvement that graph structure "
        "provides over feature-only models. The MLP baseline achieves only 0.5888 accuracy "
        "despite having access to 1,433-dimensional features, while all GNN variants "
        "exceed 0.81. This +24% absolute improvement (or +41% relative) underscores that "
        "in citation networks, the pattern of who cites whom is at least as informative "
        "as the content of the papers themselves for determining research topics."
    )
    doc.add_paragraph(
        "This result has practical implications: when building classification systems "
        "for networked data, investing in graph-based approaches yields substantially "
        "better results than purely feature-based methods, even when features are "
        "high-dimensional."
    )

    doc.add_heading('6.2 GCN vs GAT: Simplicity Wins on Small Graphs', level=2)
    doc.add_paragraph(
        "The near-identical performance of GCN (0.8276) and GAT (0.8272) challenges "
        "the assumption that more complex attention mechanisms always improve performance. "
        "On Cora's small, homogeneous graph, the simple mean aggregation of GCN captures "
        "neighborhood information as effectively as learned attention weights."
    )
    doc.add_paragraph(
        "This finding aligns with recent work (Errica et al., 2020; Shchur et al., 2018) "
        "showing that with proper hyperparameter tuning, simple GNN architectures can match "
        "or exceed more complex ones. The attention mechanism in GAT may be more beneficial "
        "on heterogeneous graphs where different neighbors have vastly different importance, "
        "or on larger graphs where the neighborhood is too large for uniform aggregation."
    )

    doc.add_heading('6.3 Residual Skip Connection: The Key Innovation', level=2)
    doc.add_paragraph(
        "The residual skip connection H₀·W_skip is the most impactful component of our "
        "proposed model, contributing +1.7% test accuracy. This connection acts as a "
        "\"shortcut\" that preserves raw feature information at the output level. Intuitively, "
        "it allows the model to combine graph-based reasoning (from the GAT layers) with "
        "direct feature-based classification (from the linear projection)."
    )
    doc.add_paragraph(
        "Interestingly, the skip connection operates on raw features H₀ rather than "
        "intermediate representations H₁. This design choice means the skip path carries "
        "the full 1,433-dimensional feature vector (projected to 7 classes), preserving "
        "fine-grained information that may be lost through the message-passing bottleneck "
        "(where features are compressed to 64 dimensions in H₁)."
    )

    doc.add_heading('6.4 DropEdge: Not Universally Beneficial', level=2)
    doc.add_paragraph(
        "Our ablation shows that DropEdge provides negligible benefit (−0.1%) on Cora. "
        "This contrasts with the original DropEdge paper (Rong et al., 2020) which "
        "demonstrated significant improvements on deeper GNNs and larger datasets. "
        "We attribute this to three factors:"
    )
    factors = [
        "Shallow architecture: Our 2-layer model is not prone to severe over-smoothing, "
        "which is the primary problem DropEdge addresses.",
        "Small graph: With only 5,278 edges, removing 20% leaves ~4,222 edges — still "
        "sufficient for effective message passing on Cora's dense clusters.",
        "Low overfitting: The val-test gap of ~1-2% indicates that the model does not "
        "overfit severely, reducing the need for regularization.",
    ]
    for f in factors:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('6.5 Training Stability', level=2)
    doc.add_paragraph(
        "All models exhibit low standard deviation across 5 random seeds (0.005-0.014), "
        "indicating that the training process is robust and not overly sensitive to "
        "initialization. The proposed model's std of 0.008 with the residual connection "
        "is comparable to other GNN baselines, suggesting that the skip connection does "
        "not introduce training instability."
    )

    # ================================================================
    # 7. LIMITATIONS AND FUTURE WORK
    # ================================================================
    doc.add_heading('7. Limitations and Future Work', level=1)

    doc.add_heading('7.1 Limitations', level=2)

    limitations = [
        ("Single Dataset", "Our experiments are limited to Cora, a small citation network "
         "with 2,708 nodes and 7 classes. The generalizability of our findings to larger, "
         "more diverse datasets (Citeseer, Pubmed, OGB) remains unverified. The effectiveness "
         "of the residual skip connection and DropEdge may differ significantly on graphs "
         "with different properties (size, density, homophily)."),
        ("Shallow Architecture", "We only evaluate 2-layer GNNs. Deeper architectures "
         "(3-5 layers) may behave differently: over-smoothing becomes more severe, "
         "potentially making DropEdge and residual connections more valuable. Our findings "
         "about DropEdge's ineffectiveness may not hold for deeper models."),
        ("Limited Regularization Study", "We only investigate DropEdge as a regularization "
         "technique. Other approaches such as Mixup on graphs, Label Smoothing, or "
         "graph-based data augmentation are not explored."),
        ("Single Metric", "We report only classification accuracy. Per-class metrics "
         "(F1-score, precision, recall), calibration analysis, and uncertainty estimation "
         "would provide a more comprehensive evaluation."),
        ("No Hyperparameter Search", "We use fixed hyperparameters based on common defaults "
         "rather than systematic tuning. A grid search or Bayesian optimization over learning "
         "rate, weight decay, dropout rate, and DropEdge probability could reveal better "
         "configurations."),
    ]
    for title, desc in limitations:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading('7.2 Future Work', level=2)

    future = [
        ("Multi-Dataset Evaluation", "Extend experiments to Citeseer, Pubmed, Arxiv-OGC, "
         "and OGB datasets to verify generalizability. Larger and more diverse graphs may "
         "reveal conditions where DropEdge and attention become more impactful."),
        ("Deeper Architectures", "Explore 3-5 layer GNNs with Jumping Knowledge connections "
         "and graph normalization layers. Deeper models would benefit more from residual "
         "connections and regularization techniques."),
        ("Graph Transformers", "Investigate Graph Transformer architectures that apply "
         "self-attention over the entire graph rather than local neighborhoods, potentially "
         "capturing long-range dependencies that 2-hop GNNs miss."),
        ("Attention Visualization", "Visualize learned attention weights to interpret what "
         "the model has learned. Understanding which edges and nodes receive the most "
         "attention could provide insights into the citation patterns."),
        ("Real-World Applications", "Apply the proposed model to practical problems such as "
         "recommendation systems, social network analysis, and molecular property prediction "
         "where graph structure is critical."),
    ]
    for title, desc in future:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # ================================================================
    # 8. CONCLUSION
    # ================================================================
    doc.add_heading('8. Conclusion', level=1)
    doc.add_paragraph(
        "This project presents a comprehensive study of Graph Neural Networks for "
        "semi-supervised node classification on the Cora citation network. We implemented "
        "and compared four models — MLP, GCN, GraphSAGE, and GAT — and proposed a "
        "Residual Multi-Head Mixing GAT that combines GATv2 multi-head attention with "
        "an input-feature skip connection and DropEdge regularization."
    )
    doc.add_paragraph(
        "Our key findings are:"
    )
    findings = [
        "Graph structure is critical: GNNs outperform MLP by +24% absolute accuracy, "
        "confirming that citation patterns carry essential information beyond paper content.",
        "The proposed model achieves the best test accuracy (0.8326 ± 0.0080), "
        "outperforming all baselines including GAT (0.8272) and GCN (0.8276).",
        "The residual skip connection is the primary contributor (+1.7%), providing "
        "a direct path from raw features to output logits.",
        "DropEdge has negligible effect on this small graph, suggesting that "
        "regularization benefits depend on graph size and model depth.",
        "All models show stable training (std 0.005-0.014) with consistent "
        "validation-based model selection.",
    ]
    for i, f in enumerate(findings, 1):
        doc.add_paragraph(f"{i}. {f}", style='List Number')

    doc.add_paragraph(
        "While our results are promising, they are limited to a single small dataset. "
        "Future work should extend evaluation to larger and more diverse graphs, explore "
        "deeper architectures, and investigate attention visualization for model "
        "interpretability. The residual skip connection concept, however, shows clear "
        "potential as a simple yet effective enhancement for graph attention networks."
    )

    # ================================================================
    # REFERENCES
    # ================================================================
    doc.add_heading('References', level=1)
    refs = [
        "Kipf, T. N., & Welling, M. (2017). Semi-Supervised Classification with Graph Convolutional Networks. ICLR.",
        "Veličković, P., et al. (2018). Graph Attention Networks. ICLR.",
        "Hamilton, W. L., Ying, R., & Leskovec, J. (2017). Inductive Representation Learning on Large Graphs. NeurIPS.",
        "Brody, S., Alon, U., & Yahav, E. (2022). How Attentive are Graph Attention Networks? ICLR.",
        "Rong, Y., et al. (2020). DropEdge: Towards Deep Graph Convolutional Networks on Node Classification. ICLR.",
        "Yang, Z., Cohen, W., & Salakhutdinov, R. (2016). Revisiting Semi-Supervised Learning with Graph Embeddings. ICML.",
        "He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep Residual Learning for Image Recognition. CVPR.",
        "Bruna, J., Zaremba, W., Szlam, A., & LeCun, Y. (2014). Spectral Networks and Deep Locally Connected Networks on Graphs. ICLR.",
        "Shchur, O., Mumme, M., Bojchevski, A., & Günnemann, S. (2018). Pitfalls of Graph Neural Network Evaluation. NeurIPS Workshop.",
        "Errica, F., Podda, M., Bacciu, D., & Micheli, A. (2020). A Fair Comparison of Graph Neural Networks for Graph Classification. ICLR.",
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f"[{i}] {ref}", style='List Number')

    # ================================================================
    # Save
    # ================================================================
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build_report()
